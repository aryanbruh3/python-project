from docx import Document
import os
filename="record.docx"

def write_record():
    doc=Document()
    doc.add_heading("       STUDENT RECORDS       ",level=1)
    table=doc.add_table(rows=1,cols=3)
    table.style="Table Grid"

    hdr_cells=table.rows[0].cells
    hdr_cells[0].text="ID"
    hdr_cells[1].text="Name"
    hdr_cells[2].text="Age"

    print("\nEnter records (type 'stop') as ID to finish")
    while True:
        new_id=input("enter ID = ")
        if new_id=="stop":
            break
        else:
            new_name=input("enter Name = ")
            new_age=input("enter Age = ")
            row_cells=table.add_row().cells
            row_cells[0].text, row_cells[1].text, row_cells[2].text = new_id,new_name,new_age
    doc.save(filename)
    print("record written successfully",filename)


def read_record():
    if not os.path.exists(filename):
        print("no file found")
        return
    doc=Document(filename)
    table=doc.tables[0]
    print("\nall records")
    for row in table.rows[1:]:
        print(row.cells[0].text, row.cells[1].text, row.cells[2].text)

def append_record():
    if os.path.exists(filename):
        doc=Document(filename)
        table=doc.tables[0]
        print("\nEnter records (type 'stop') as ID to finish")
    while True:
        new_id=input("enter ID = ")
        if new_id=="stop":
            break
        else:
            new_name=input("enter Name = ")
            new_age=input("enter Age = ")
            row_cells=table.add_row().cells
            row_cells[0].text, row_cells[1].text, row_cells[2].text = new_id,new_name,new_age
        doc.save(filename)
        print("record append successfully")

def update_record():
    if not os.path.exists(filename):
        print("no record found")
        return 
    doc=Document(filename)
    table=doc.tables[0]

    search_id=input("enter the ID to update = ")
    found=False
    for row in table.rows[1:]:
        if row.cells[0].text==search_id:
            print("revord found",row.cells[0].text,row.cells[1].text,row.cells[2].text)
            print("1 - update ID\n2 - update name\n3 - update age\n4 - update all\n")
            ch=input("enter ur choice = ")
            if ch==1:
                row.cells[0].text=input("enter new ID = ")
                print("enter ID has been update successfully")
                found = True
                break

            if ch==2:
                row.cells[1].text=input("enter new name = ")
                print("name has been updated successfully")
                found=True
                break

            if ch==3:
                row.cells[2].text=input("enter new age = ")
                print("age has been updated successfully")
                found=True
                break

            if ch==4:
                row.cells[0].text=input("enter new ID = ")
                row.cells[1].text=input("enter new name = ")
                row.cells[2].text=input("enter new age = ")
                print("ID,name and age has been updated successfully")
                found=True
                break

        else:
            print("no record has been updated")
            break

    if found:
        doc.save(filename)
        print("record has been updates successfully")
    else:
        print("no record has been found")


def delete_record():
    if os.path.exists(filename):
        doc=Document(filename)
        table=doc.tables[0]
    
    search_id=input("enter id to delete = ")
    new_record=[]
    for row in table.rows[1:]:
        if row.cells[0].text==search_id:
            print("deleting record",row.cells[0].text,row.cells[1].text,row.cells[2].text)
            found=True
        else:
            new_row.append([row_cells[0].text,row_cells[1].text,row_cells[2].text])

    if found:
        doc=Document()
        doc.add_heading("   Student record   ")
        doc.add_heading("       STUDENT RECORDS       ",level=1)
        new_table=doc.add_table(rows=1,cols=3)
        new_table.style="Table Grid"
        hdr_cells=new_table.rows[0].cells
        hdr_cells[0].text,hdr_cells[1].text,hdr_cells[2].text=["ID","name","age"]

        for rec in new_rows:
            row_cells=new_table.add_row().cells
            row_cells[0].text,row_cells[1].text,row_cells[2].text=rec
        doc.save(filename)
        print("record deleted successfully")

# 1 - how many record tables data u want to store in file
# 2 - enter first table heading name
#       roll no -  
#       student name - 
#       marks - 
#     enter second table
#       emplyee id
#       name
#       salary
# 3 - in which table u want to insert recordzkl